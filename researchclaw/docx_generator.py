"""Generador de documentos Word (.docx) desde Markdown.

Detecta automáticamente si el contenido es un Informe de Auditoría CEIm
y aplica el formato estructurado correspondiente (portada, Resumen con borde,
tabla de Checklist, Dictamen con color de dictamen).

Para documentos de investigación general aplica formato académico estándar.

Uso:
    from researchclaw.docx_generator import generate_docx
    path = generate_docx(markdown_text, output_path, title="Mi informe")
"""

from __future__ import annotations

import re
from pathlib import Path
from datetime import date


# ---------------------------------------------------------------------------
# APA-7 citation formatting
# ---------------------------------------------------------------------------

def _parse_bibtex_field(entry: str, field: str) -> str:
    """Extract a single field value from a BibTeX entry string."""
    pat = re.compile(
        rf"\b{re.escape(field)}\s*=\s*\{{(.+?)\}}\s*[,}}]",
        re.IGNORECASE | re.DOTALL,
    )
    m = pat.search(entry)
    return m.group(1).strip() if m else ""


def _apa_authors(raw_authors: str) -> str:
    """Convert BibTeX author string to APA format.

    Input: ``Doe, John and Smith, Jane`` or ``John Doe and Jane Smith``
    Output: ``Doe, J., & Smith, J.``
    """
    if not raw_authors:
        return "Unknown Author"
    parts = re.split(r"\s+and\s+", raw_authors, flags=re.IGNORECASE)
    formatted: list[str] = []
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if "," in part:
            # Last, First → Last, F.
            last, _, first = part.partition(",")
            initials = "".join(
                f"{w[0].upper()}." for w in first.strip().split() if w
            )
            formatted.append(f"{last.strip()}, {initials}" if initials else last.strip())
        else:
            # First Last → Last, F.
            words = part.split()
            if len(words) == 1:
                formatted.append(words[0])
            else:
                last = words[-1]
                initials = "".join(f"{w[0].upper()}." for w in words[:-1] if w)
                formatted.append(f"{last}, {initials}" if initials else last)
    if len(formatted) == 1:
        return formatted[0]
    return ", ".join(formatted[:-1]) + ", & " + formatted[-1]


def _bibtex_entry_to_apa(entry: str) -> str:
    """Convert a single BibTeX entry to an APA-7 formatted reference string.

    Format: Author, A. A. (Year). Title. Journal, volume(issue), pages.
            https://doi.org/…
    """
    entry_type_m = re.match(r"@(\w+)\{", entry)
    entry_type = entry_type_m.group(1).lower() if entry_type_m else "article"

    authors = _parse_bibtex_field(entry, "author")
    year    = _parse_bibtex_field(entry, "year")
    title   = _parse_bibtex_field(entry, "title")
    journal = (
        _parse_bibtex_field(entry, "journal")
        or _parse_bibtex_field(entry, "booktitle")
    )
    volume  = _parse_bibtex_field(entry, "volume")
    number  = _parse_bibtex_field(entry, "number")
    pages   = _parse_bibtex_field(entry, "pages")
    doi     = _parse_bibtex_field(entry, "doi")
    url     = _parse_bibtex_field(entry, "url")
    eprint  = _parse_bibtex_field(entry, "eprint")

    apa_authors = _apa_authors(authors)
    year_str = f"({year})" if year else "(n.d.)"
    title_str = title or "(No title)"

    parts = [f"{apa_authors} {year_str}. {title_str}."]

    if entry_type == "inproceedings" and journal:
        parts.append(f" In *{journal}*.")
    elif journal and "arxiv preprint" not in journal.lower():
        journal_part = f" *{journal}*"
        if volume:
            journal_part += f", *{volume}*"
            if number:
                journal_part += f"({number})"
        if pages:
            journal_part += f", {pages}"
        parts.append(journal_part + ".")
    elif eprint:
        parts.append(f" arXiv preprint arXiv:{eprint}.")

    if doi:
        doi_clean = doi.lstrip("https://doi.org/").lstrip("http://doi.org/")
        parts.append(f" https://doi.org/{doi_clean}")
    elif url and not eprint:
        parts.append(f" {url}")

    return "".join(parts)


def bibtex_to_apa_bibliography(bib_text: str) -> list[str]:
    """Parse a BibTeX file and return list of APA-formatted reference strings.

    Returns an empty list if the input is empty or contains no valid entries.
    """
    if not bib_text or not bib_text.strip():
        return []
    # Split into individual @type{...} blocks
    entries: list[str] = []
    depth = 0
    start = -1
    for i, ch in enumerate(bib_text):
        if ch == "@" and depth == 0:
            start = i
        elif ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0 and start != -1:
                entries.append(bib_text[start : i + 1])
                start = -1
    results: list[str] = []
    for entry in entries:
        if re.match(r"@string\b|@comment\b|@preamble\b", entry, re.I):
            continue
        try:
            results.append(_bibtex_entry_to_apa(entry))
        except Exception:  # noqa: BLE001
            continue
    return results


# ---------------------------------------------------------------------------
# Utilidades de parsing Markdown
# ---------------------------------------------------------------------------

def _strip_inline_md(text: str) -> str:
    """Elimina marcadores de negrita/cursiva inline."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text.strip()


def _parse_md_table(lines: list[str]) -> list[list[str]]:
    """Parsea un bloque de tabla Markdown. Devuelve lista de filas."""
    rows: list[list[str]] = []
    for line in lines:
        if re.match(r"^\s*\|[-:| ]+\|\s*$", line):
            continue  # separador
        cells = [_strip_inline_md(c.strip()) for c in line.strip().strip("|").split("|")]
        if cells:
            rows.append(cells)
    return rows


def _classify_content(text: str) -> str:
    """Devuelve 'ceim' si parece un informe CEIm, 'general' en caso contrario."""
    ceim_markers = [
        "DICTAMEN", "Checklist de Auditoría CEIm", "Auditoría CEIm",
        "ACEPTAR", "ACLARACIONES", "RECHAZAR",
        "Ítems con ✅", "Ítem 1 ⭐", "BLOQUE A —",
    ]
    return "ceim" if any(m in text for m in ceim_markers) else "general"


# ---------------------------------------------------------------------------
# Estilos y colores
# ---------------------------------------------------------------------------

_BLUE_DARK  = "1F3864"   # azul oscuro corporativo
_BLUE_MID   = "2E74B5"   # azul medio títulos
_BLUE_LIGHT = "D6EAF8"   # fondo tabla cabecera
_GREEN      = "D9EAD3"   # fondo dictamen ACEPTAR
_ORANGE     = "FCE5CD"   # fondo dictamen ACLARACIONES
_RED        = "FCE4D6"   # fondo dictamen RECHAZAR
_GREY_LIGHT = "F2F2F2"   # filas alternas tabla


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _set_cell_shading(cell, hex_color: str) -> None:
    """Aplica color de fondo a una celda de tabla via XML."""
    try:
        from docx.oxml.ns import qn  # type: ignore[import]
        from docx.oxml import OxmlElement  # type: ignore[import]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)
    except Exception:
        pass


def _set_paragraph_border(para, hex_color: str = "2E74B5") -> None:
    """Añade borde izquierdo de acento a un párrafo."""
    try:
        from docx.oxml.ns import qn  # type: ignore[import]
        from docx.oxml import OxmlElement  # type: ignore[import]
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "24")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), hex_color)
        pBdr.append(left)
        pPr.append(pBdr)
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Helpers de Document
# ---------------------------------------------------------------------------

def _add_heading_styled(doc, text: str, level: int, color_hex: str = _BLUE_MID) -> None:
    from docx.shared import RGBColor, Pt  # type: ignore[import]
    heading = doc.add_heading(_strip_inline_md(text), level=level)
    heading.runs[0].font.color.rgb = RGBColor(*_hex_to_rgb(color_hex))
    if level == 1:
        heading.runs[0].font.size = Pt(18)
    elif level == 2:
        heading.runs[0].font.size = Pt(14)
    else:
        heading.runs[0].font.size = Pt(12)


def _add_paragraph_with_inline_bold(doc, text: str, style: str = "Normal") -> None:
    """Añade un párrafo respetando el **bold** inline."""
    from docx.shared import Pt  # type: ignore[import]
    para = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)
    return para


def _add_table_from_rows(doc, rows: list[list[str]], is_ceim: bool = False):
    """Crea una tabla Word a partir de filas de texto."""
    if not rows:
        return
    from docx.shared import Pt, RGBColor  # type: ignore[import]

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data[:n_cols]):
            cell = row.cells[j]
            cell.text = cell_text
            para = cell.paragraphs[0]
            para.runs[0].font.size = Pt(9) if is_ceim else Pt(10)

            # Cabecera → fondo azul + negrita
            if i == 0:
                _set_cell_shading(cell, _BLUE_LIGHT)
                para.runs[0].bold = True
                para.runs[0].font.color.rgb = RGBColor(*_hex_to_rgb(_BLUE_DARK))
            # Filas alternas
            elif i % 2 == 0:
                _set_cell_shading(cell, _GREY_LIGHT)

            # Para CEIm: colorear estado ✅/⚠️/❌
            if is_ceim and i > 0 and j == 1:
                emoji = cell_text.strip()
                if "✅" in emoji:
                    _set_cell_shading(cell, "D9EAD3")
                elif "⚠️" in emoji or "⚠" in emoji:
                    _set_cell_shading(cell, "FFF2CC")
                elif "❌" in emoji:
                    _set_cell_shading(cell, "FCE4D6")

    doc.add_paragraph()  # espaciado tras tabla


# ---------------------------------------------------------------------------
# Estilo global del documento
# ---------------------------------------------------------------------------

def _set_doc_style(doc) -> None:
    """Establece Arial 11pt, interlineado 1.5 y espacio tras párrafo en el estilo Normal."""
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING  # type: ignore[import]

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_after = Pt(4)

    # Propagar a List Bullet si existe
    try:
        lb = doc.styles["List Bullet"]
        lb.font.name = "Arial"
        lb.font.size = Pt(11)
        lb.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Pie de página — 3 columnas con numeración automática
# ---------------------------------------------------------------------------

def _add_word_field(para, field_code: str, font_size_pt: int = 9) -> None:
    """Inserta un campo Word (PAGE / NUMPAGES) en el párrafo dado."""
    from docx.oxml import OxmlElement  # type: ignore[import]
    from docx.oxml.ns import qn         # type: ignore[import]
    from docx.shared import Pt, RGBColor  # type: ignore[import]

    r = para.add_run()
    r.font.size = Pt(font_size_pt)
    r.font.color.rgb = RGBColor(130, 130, 130)

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    r._r.append(fldChar_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    r._r.append(instr)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    r._r.append(fldChar_end)


def _add_docx_footer(doc, short_title: str, version: str) -> None:
    """Pie de página tripartito: título corto | versión · fecha | Página X de Y."""
    from docx.oxml import OxmlElement  # type: ignore[import]
    from docx.oxml.ns import qn         # type: ignore[import]
    from docx.shared import Pt, RGBColor  # type: ignore[import]

    today = date.today().strftime("%d/%m/%Y")
    # Posiciones de tabulación en twips (1 cm ≈ 567 twips)
    # A4 usable width con márgenes 2.5 cm: 21 cm - 5 cm = 16 cm = 9072 twips
    TAB_CENTER = "4536"   # 8 cm → centro
    TAB_RIGHT  = "9072"   # 16 cm → margen derecho

    def _grey_run(para, text: str) -> None:
        r = para.add_run(text)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(130, 130, 130)

    def _tab_run(para) -> None:
        from docx.oxml import OxmlElement as OE
        r = para.add_run()
        tab = OE("w:tab")
        r._r.append(tab)

    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        # Usar (o limpiar) el primer párrafo del footer
        if footer.paragraphs:
            fp = footer.paragraphs[0]
            fp.clear()
        else:
            fp = footer.add_paragraph()

        # Definir tab stops en el párrafo
        pPr = fp._p.get_or_add_pPr()
        tabs_elem = OxmlElement("w:tabs")
        for val, pos in [("center", TAB_CENTER), ("right", TAB_RIGHT)]:
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), val)
            tab.set(qn("w:pos"), pos)
            tabs_elem.append(tab)
        pPr.append(tabs_elem)

        # Columna izquierda: título corto
        _grey_run(fp, short_title[:50])

        # Separador → columna central
        _tab_run(fp)
        _grey_run(fp, f"{version}  ·  {today}")

        # Separador → columna derecha: Página X de Y
        _tab_run(fp)
        _grey_run(fp, "Página ")
        _add_word_field(fp, "PAGE")
        _grey_run(fp, " de ")
        _add_word_field(fp, "NUMPAGES")


# ---------------------------------------------------------------------------
# Portada
# ---------------------------------------------------------------------------

def _add_cover_page(doc, title: str, doc_type: str) -> None:
    from docx.shared import Pt, RGBColor  # type: ignore[import]
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]

    # Espacio superior
    for _ in range(6):
        doc.add_paragraph()

    # Título principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(*_hex_to_rgb(_BLUE_DARK))

    # Subtítulo tipo
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p_sub.add_run(doc_type)
    run2.font.size = Pt(13)
    run2.font.color.rgb = RGBColor(*_hex_to_rgb(_BLUE_MID))

    doc.add_paragraph()

    # Fecha
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p_date.add_run(f"Fecha: {date.today().strftime('%d de %B de %Y')}")
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(100, 100, 100)

    # Pie de portada
    for _ in range(4):
        doc.add_paragraph()
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p_foot.add_run("Generado por ResearchClaw — IA Médica Autónoma")
    run4.font.size = Pt(9)
    run4.font.color.rgb = RGBColor(150, 150, 150)
    run4.italic = True

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Generador CEIm específico
# ---------------------------------------------------------------------------

def _generate_ceim_docx(doc, sections: list[tuple[str, str]]) -> None:
    """Formatea el informe CEIm con sus secciones propias."""
    from docx.shared import Pt, RGBColor  # type: ignore[import]
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]

    for heading, body in sections:
        heading_clean = _strip_inline_md(heading)

        # Detectar sección Dictamen para darle color especial
        is_dictamen = re.search(r"dictamen|verdict", heading_clean, re.I)

        if heading_clean:
            _add_heading_styled(doc, heading_clean, level=2)

        # Detectar si el cuerpo contiene una tabla Markdown
        table_lines = [l for l in body.splitlines() if l.strip().startswith("|")]
        if len(table_lines) >= 2:
            rows = _parse_md_table(table_lines)
            if rows:
                _add_table_from_rows(doc, rows, is_ceim=True)
            # Texto antes/después de la tabla
            non_table = "\n".join(
                l for l in body.splitlines() if not l.strip().startswith("|")
            ).strip()
            if non_table:
                _add_paragraph_with_inline_bold(doc, non_table)
        else:
            # Procesar párrafos y bullets
            for line in body.splitlines():
                line = line.rstrip()
                if not line:
                    continue
                if line.lstrip().startswith(("- ", "* ")):
                    bullet_text = _strip_inline_md(line.lstrip("- *").strip())
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(bullet_text)
                else:
                    para = _add_paragraph_with_inline_bold(doc, line)
                    if is_dictamen:
                        # Colorear el fondo del dictamen
                        text_upper = line.upper()
                        if "ACEPTAR" in text_upper and "RECHAZAR" not in text_upper:
                            _set_paragraph_border(para, "2EA02E")
                        elif "RECHAZAR" in text_upper:
                            _set_paragraph_border(para, "C00000")
                        elif "ACLARACIONES" in text_upper:
                            _set_paragraph_border(para, "ED7D31")

        doc.add_paragraph()  # espacio entre secciones


# ---------------------------------------------------------------------------
# Generador general (investigación / artículo)
# ---------------------------------------------------------------------------

_REFERENCES_HEADINGS = frozenset({
    "references", "referencias", "bibliography", "bibliografía",
    "referencias bibliográficas", "works cited", "literature cited",
})


def _generate_general_docx(
    doc,
    sections: list[tuple[str, str]],
    bib_text: str = "",
) -> None:
    from docx.shared import Pt  # type: ignore[import]
    from docx.enum.text import WD_LINE_SPACING  # type: ignore[import]

    # Pre-compute APA references if bib_text provided
    apa_refs: list[str] = bibtex_to_apa_bibliography(bib_text) if bib_text else []

    for heading, body in sections:
        heading_clean = _strip_inline_md(heading)
        if heading_clean:
            _add_heading_styled(doc, heading_clean, level=2)

        # Detect References/Bibliography section → render in APA format
        is_refs_section = heading_clean.strip().lower() in _REFERENCES_HEADINGS
        if is_refs_section and apa_refs:
            for ref in apa_refs:
                # APA references use hanging indent (1.27 cm) per APA-7 spec
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(28)       # ~1 cm
                p.paragraph_format.first_line_indent = Pt(-28)  # hanging
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                run = p.add_run(ref)
                run.font.size = Pt(10)
            doc.add_paragraph()
            continue  # skip body rendering for this section

        table_lines = [l for l in body.splitlines() if l.strip().startswith("|")]
        if len(table_lines) >= 2:
            rows = _parse_md_table(table_lines)
            if rows:
                _add_table_from_rows(doc, rows, is_ceim=False)

        for line in body.splitlines():
            line = line.rstrip()
            if not line or (line.strip().startswith("|")):
                continue
            if line.lstrip().startswith(("- ", "* ")):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(_strip_inline_md(line.lstrip("- *").strip()))
            else:
                _add_paragraph_with_inline_bold(doc, line)

        doc.add_paragraph()


# ---------------------------------------------------------------------------
# Función pública principal
# ---------------------------------------------------------------------------

def generate_docx(
    markdown_text: str,
    output_path: Path,
    title: str = "Informe ResearchClaw",
    version: str = "Versión 1.0",
    short_title: str = "",
    bib_text: str = "",
) -> Path:
    """Genera un fichero .docx a partir del Markdown del paper/informe.

    Detecta automáticamente si es un informe CEIm y aplica el formato
    estructurado con portada, secciones coloreadas y tabla de checklist.

    Parameters
    ----------
    version:
        Texto de versión que aparecerá en el pie (ej. "Versión 1.0").
    short_title:
        Título corto o código de protocolo para el pie izquierdo.
        Si se omite se usa los primeros 40 caracteres del título.
    bib_text:
        Contenido BibTeX completo (de references.bib).  Cuando se proporciona,
        la sección References se formatea automáticamente en estilo APA-7.
    """
    try:
        from docx import Document  # type: ignore[import]
        from docx.shared import Pt, Cm  # type: ignore[import]
        from docx.oxml.ns import qn  # type: ignore[import]
    except ImportError as exc:
        raise ImportError("python-docx no instalado. Ejecuta: pip install python-docx") from exc

    content_type = _classify_content(markdown_text)

    # Extraer título del H1 si existe
    h1 = re.search(r"^#\s+(.+)$", markdown_text, re.MULTILINE)
    doc_title = h1.group(1).strip() if h1 else title
    footer_left = short_title.strip() if short_title.strip() else doc_title[:40]

    # Parsear secciones ## y ###
    sections: list[tuple[str, str]] = []
    pattern = re.compile(r"^#{2,3}\s+(.+)$", re.MULTILINE)
    matches = list(pattern.finditer(markdown_text))

    # Texto antes del primer ## (intro/resumen ejecutivo)
    preamble = markdown_text[: matches[0].start()].strip() if matches else markdown_text

    for i, m in enumerate(matches):
        heading = m.group(1).strip()
        start   = m.end()
        end     = matches[i + 1].start() if i + 1 < len(matches) else len(markdown_text)
        body    = markdown_text[start:end].strip()
        sections.append((heading, body))

    doc = Document()

    # Estilo global: Arial 11pt, interlineado 1.5
    _set_doc_style(doc)

    # Márgenes estándar 2.5 cm
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Portada
    doc_type_label = (
        "Informe de Auditoría Ética — CEIm" if content_type == "ceim"
        else "Informe de Investigación Científica"
    )
    _add_cover_page(doc, doc_title, doc_type_label)

    # Preámbulo (resumen ejecutivo / abstract del paper)
    if preamble:
        preamble_clean = re.sub(r"^#\s+.+$", "", preamble, flags=re.MULTILINE).strip()
        if preamble_clean:
            _add_heading_styled(doc, "Resumen Ejecutivo", level=1, color_hex=_BLUE_DARK)
            for line in preamble_clean.splitlines():
                if line.strip():
                    _add_paragraph_with_inline_bold(doc, line.strip())
            doc.add_paragraph()

    # Secciones
    if content_type == "ceim":
        _generate_ceim_docx(doc, sections)
    else:
        _generate_general_docx(doc, sections, bib_text=bib_text)

    # Pie de página tripartito
    _add_docx_footer(doc, footer_left, version)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
